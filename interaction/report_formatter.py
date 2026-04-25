#!/usr/bin/env python3
"""Report Formatter: generates comprehensive SEM analysis report with audit trail."""

from datetime import datetime
from typing import Dict, Any, List, Optional
import os
import pandas as pd

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from utils.state_manager import SessionState


class ReportFormatter:
    """Generates Markdown reports from session state."""

    def __init__(self):
        pass

    def generate(self, session: SessionState, output_path: Optional[str] = None, format: str = 'md') -> str:
        """Generate full report. Returns file path if saved, else markdown string.
        Args:
            session: SessionState object
            output_path: file path to save (extension determines format if format not given)
            format: 'md' or 'docx'
        """
        md = self._build_report(session)

        if output_path:
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
            if format == 'docx' or output_path.endswith('.docx'):
                if not DOCX_AVAILABLE:
                    return f"python-docx not installed. Cannot generate Word document. Run: pip install python-docx\nMarkdown version generated instead: {output_path.replace('.docx', '.md')}"
                docx_path = output_path if format == 'docx' else output_path
                self._save_as_docx(md, docx_path)
                return docx_path
            else:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(md)
                return output_path
        return md

    def _build_report(self, session: SessionState) -> str:
        lines = []

        # Title block
        lines.append("# Structural Equation Modeling (SEM) Analysis Report")
        lines.append(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        lines.append("")

        # 1. 研究设计与数据
        lines.append("## 1. Research Design and Data")
        df = session.data
        lines.append(f"- **Sample size**: {len(df)}")
        lines.append(f"- **Total variables**: {len(df.columns)}")
        lines.append(f"- **Variable list**: {', '.join(df.columns)}")
        lines.append("")

        # Hypotheses
        raw_hyp = session.variables.get('raw_hypothesis', '')
        if raw_hyp:
            lines.append("### Research Hypothesis (Original Description)")
            lines.append(f"> {raw_hyp}")
            lines.append("")

        # Structured hypotheses
        paths = getattr(session, 'suggested_paths', [])
        if paths:
            lines.append("### Structured Hypotheses (Paths)")
            for src, dst in paths:
                lines.append(f"- {src} → {dst}")
            lines.append("")

        # 2. 变量定义与测量模型
        lines.append("## 2. Variable Definitions and Measurement Model")
        latent_inds = session.variables.get('latent_indicators', {})
        if latent_inds:
            lines.append("### Latent Variables")
            for lv, items in latent_inds.items():
                lines.append(f"- **{lv}** =~ {', '.join(items)}")
            lines.append("")
        classifications = session.variables.get('classifications', {})
        if classifications:
            lines.append("### Observed Variables (Directly Measured)")
            observed = [v for v, t in classifications.items() if t == 'observed']
            if observed:
                lines.append(f"- {', '.join(observed)}")
            lines.append("")
        remaining_class = getattr(session, 'remaining_classification', {})
        covariates = [v for v, r in remaining_class.items() if r == 'covariate']
        if covariates:
            lines.append("### Covariates")
            lines.append(f"- {', '.join(covariates)}")
            lines.append("")

        # 3. 分析方法与决策
        lines.append("## 3. Analysis Method and Decisions")
        lines.append(f"- **Missing data handling**: {getattr(session, 'missing_handling', 'Not specified')}")
        lines.append("- **Estimation method**: ML (Maximum Likelihood)")
        lines.append("")

        # Decision log (beautified)
        decision_log = getattr(session, 'decision_log', [])
        if decision_log:
            lines.append("### Analysis Log (Audit Trail)")
            lines.append("| Time | Event | Details |")
            lines.append("|------|------|------|")
            for entry in decision_log:
                ts = entry.get('timestamp', '')
                event = entry.get('event', '')
                detail = entry.get('detail', '')
                if isinstance(detail, (dict, list)):
                    import json
                    detail_str = json.dumps(detail, ensure_ascii=False)
                else:
                    detail_str = str(detail)
                # Truncate long details
                if len(detail_str) > 80:
                    detail_str = detail_str[:77] + "..."
                lines.append(f"| {ts} | {event} | {detail_str} |")
            lines.append("")
        # Also include full session history if decision_log is sparse
        history = getattr(session, 'history', [])
        if history and len(decision_log) < len(history):
            lines.append("### Full Operation Log")
            for entry in history[-20:]:  # last 20 entries
                ts = entry.get('timestamp', '')
                action = entry.get('action', '')
                details = entry.get('details', '')
                if isinstance(details, (dict, list)):
                    import json
                    details_str = json.dumps(details, ensure_ascii=False)
                else:
                    details_str = str(details)
                lines.append(f"- {ts}: {action} - {details_str}")
            lines.append("")

        # 4. 模型结果
        results = getattr(session, 'results', None)
        if results:
            lines.append("## 4. Model Fit Results")
            fit = results.get('fit_indices', {})
            if fit and 'error' not in fit:
                lines.append("### Fit Indices")
                lines.append("| Index | Value |")
                lines.append("|------|-----|")
                for key, label in [
                    ('chi2', 'χ²'),
                    ('df', 'df'),
                    ('p_value', 'p-value'),
                    ('cfi', 'CFI'),
                    ('tli', 'TLI'),
                    ('rmsea', 'RMSEA'),
                    ('srmr', 'SRMR'),
                    ('aic', 'AIC'),
                    ('bic', 'BIC')
                ]:
                    val = fit.get(key)
                    if val is not None:
                        if isinstance(val, float):
                            val_str = f"{val:.3f}"
                        else:
                            val_str = str(val)
                        lines.append(f"| {label} | {val_str} |")
                lines.append("")

            paths_res = results.get('paths', {})
            if paths_res and not paths_res.get('error'):
                lines.append("### Path Coefficients (Standardized)")
                lines.append("| Path | Estimate | Std. Error | z | p | Significance |")
                lines.append("|------|------|--------|-----|------|--------|")
                for (src, dst), stats in paths_res.items():
                    est = stats.get('estimate')
                    se = stats.get('se')
                    z = stats.get('z')
                    p = stats.get('p')
                    sig = '***' if p is not None and p < 0.001 else ('**' if p is not None and p < 0.01 else ('*' if p is not None and p < 0.05 else ''))
                    est_str = f"{est:.3f}" if est is not None else "N/A"
                    se_str = f"{se:.3f}" if se is not None else "N/A"
                    z_str = f"{z:.3f}" if z is not None else "N/A"
                    p_str = f"{p:.3f}" if p is not None else "N/A"
                    lines.append(f"| {src} → {dst} | {est_str} | {se_str} | {z_str} | {p_str} | {sig} |")
                lines.append("")

        # 5. 高级分析结果（如果有）
        advanced_results = getattr(session, 'advanced_results', {})
        if advanced_results:
            lines.append("## 5. Advanced Analysis Results")
            for model_type, result in advanced_results.items():
                lines.append(f"### {model_type.upper()}")
                fit = result.fit_stats
                lines.append(f"- CFI = {fit.get('cfi', 'N/A')}, TLI = {fit.get('tli', 'N/A')}")
                lines.append(f"- RMSEA = {fit.get('rmsea', 'N/A')}, SRMR = {fit.get('srmr', 'N/A')}")
                lines.append(f"- χ²({fit.get('df','?')}) = {fit.get('chi2',0):.2f}, p = {fit.get('p_value','N/A')}")
                lines.append("")
                # Parameters table
                params = result.parameters
                if isinstance(params, pd.DataFrame):
                    lines.append("| Path/Parameter | Estimate | Std. Error | z | p |")
                    lines.append("|-----------|------|--------|---|----|")
                    for _, row in params.head(15).iterrows():
                        path = row.get('path', '?')
                        est = row.get('estimate', None)
                        se = row.get('se', None)
                        z = row.get('z', None)
                        p = row.get('p', None)
                        est_str = f"{est:.3f}" if isinstance(est, (int,float)) else str(est)
                        se_str = f"{se:.3f}" if isinstance(se, (int,float)) else str(se)
                        z_str = f"{z:.3f}" if isinstance(z, (int,float)) else str(z)
                        p_str = f"{p:.3f}" if isinstance(p, (int,float)) else str(p)
                        lines.append(f"| {path} | {est_str} | {se_str} | {z_str} | {p_str} |")
                    lines.append("")
                if result.bootstrap_ci:
                    lines.append("**Bootstrap 95% Confidence Intervals**:")
                    for path, (lo, up) in result.bootstrap_ci.items():
                        lines.append(f"- {path}: [{lo:.3f}, {up:.3f}]")
                    lines.append("")

        # 6. Recommended References
        lines.append("## 6. Recommended References")
        retriever = getattr(session, 'textbook_retriever', None)
        if retriever:
            stages_used = ['measurement_model', 'structural_model', 'data_preparation',
                           'model_specification', 'reporting', 'model_modification']
            seen_ids = set()
            for stage in stages_used:
                try:
                    sections = retriever.get_by_stage(stage, top_k=2)
                    for sec in sections:
                        key = f"{sec.title}|{sec.source_short}"
                        if key not in seen_ids:
                            seen_ids.add(key)
                            lines.append(f"- {sec.title} ({sec.source_short})")
                except Exception:
                    pass
            if not seen_ids:
                lines.append("- No textbook references available")
        else:
            lines.append("- Textbook retrieval not available")
        lines.append("")

        # Appendix
        lines.append("## Appendix")
        lines.append(f"- Analysis tool: sem-analyzer skill (OpenClaw)")
        lines.append(f"- Python environment: semopy {self._get_semopy_version()}")
        # Add model specification (semopy syntax)
        draft = getattr(session, 'draft', None)
        if draft:
            try:
                model_spec = draft.to_semopy_model()
                lines.append("### Model Definition (semopy syntax)")
                lines.append("```model\n" + model_spec + "\n```")
            except Exception:
                pass
        lines.append("")

        return "\n".join(lines)

    def _get_semopy_version(self) -> str:
        try:
            import semopy
            return getattr(semopy, '__version__', 'unknown')
        except:
            return 'unavailable'

    def _save_as_docx(self, markdown: str, output_path: str):
        """Convert simple markdown to DOCX using python-docx.
        Supports: headings (#, ##, ###), bold (**text**), italic (*text*), lists (-, *), tables (pipe).
        """
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()

        lines = markdown.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                i += 1
                continue
            # Headings
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                i += 1
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                i += 1
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                i += 1
                continue
            # Table detection
            if line.startswith('|') and line.endswith('|') and '|' in line[1:-1]:
                # Collect header, separator, rows
                header = [c.strip() for c in line.strip('|').split('|')]
                i += 1
                if i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|') and all('---' in c or ':-:' in c for c in lines[i].strip('|').split('|')):
                    i += 1  # skip separator
                rows = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    row = [c.strip() for c in lines[i].strip('|').split('|')]
                    rows.append(row)
                    i += 1
                # Add table
                if rows:
                    table = doc.add_table(rows=len(rows)+1, cols=len(header))
                    table.style = 'Light Shading'
                    # Header
                    for j, h in enumerate(header):
                        table.cell(0, j).text = h
                        table.cell(0, j).paragraphs[0].runs[0].bold = True
                    # Data rows
                    for r_i, row in enumerate(rows):
                        for j, cell in enumerate(row):
                            table.cell(r_i+1, j).text = cell
                continue
            # List items
            if line.startswith('- ') or line.startswith('* '):
                bullet = True
                items = []
                while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                    items.append(lines[i][2:])
                    i += 1
                # Add as bullet list
                for it in items:
                    doc.add_paragraph(it, style='List Bullet')
                continue
            # Plain paragraph with inline formatting
            if line:
                p = doc.add_paragraph()
                # Simple bold/italic handling: **bold** and *italic*
                # We'll split by ** and *
                parts = []
                # Bold **
                tokens = line.split('**')
                odd = False
                for token in tokens:
                    if odd:
                        parts.append(('bold', token))
                    else:
                        parts.append(('normal', token))
                    odd = not odd
                # Italic *
                new_parts = []
                for typ, txt in parts:
                    if typ == 'normal' and '*' in txt:
                        sub = txt.split('*')
                        sub_odd = False
                        for s in sub:
                            if sub_odd:
                                new_parts.append(('italic', s))
                            else:
                                new_parts.append(('normal', s))
                            sub_odd = not sub_odd
                    else:
                        new_parts.append((typ, txt))
                # Add runs
                for typ, txt in new_parts:
                    run = p.add_run(txt)
                    if typ == 'bold':
                        run.bold = True
                    if typ == 'italic':
                        run.italic = True
                i += 1
                continue
            i += 1

        doc.save(output_path)
